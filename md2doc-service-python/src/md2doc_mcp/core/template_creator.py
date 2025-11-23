"""
动态 Word 文档模板创建器
对应 Java 版本的 DynamicWordDocumentCreator.java
"""

import re
from typing import List, Tuple
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from ..parser.markdown_parser import MarkdownParser


def format_number_with_thousands_separator(text: str) -> str:
    """格式化文本中的数字，添加千分符号
    
    例如：1000 -> 1,000，1234567 -> 1,234,567
    
    Args:
        text: 输入文本
        
    Returns:
        格式化后的文本
    """
    # 匹配整数和小数的正则表达式
    def replace_number(match):
        number_str = match.group(0)
        # 如果是小数，分别处理整数部分和小数部分
        if '.' in number_str:
            integer_part, decimal_part = number_str.split('.')
            # 格式化整数部分
            formatted_integer = format(int(integer_part), ',')
            return f"{formatted_integer}.{decimal_part}"
        else:
            # 格式化整数
            return format(int(number_str), ',')
    
    # 匹配数字（包括整数和小数）
    pattern = r'\d+(\.\d+)?'
    return re.sub(pattern, replace_number, text)


def number_to_chinese(num: int) -> str:
    """将数字转换为中文数字
    
    Args:
        num: 数字 (1-99)
        
    Returns:
        中文数字字符串，如 "一"、"二"、"十"、"十一" 等
    """
    if num <= 0 or num > 99:
        return str(num)
    
    # 基本数字映射
    chinese_digits = ['', '一', '二', '三', '四', '五', '六', '七', '八', '九']
    
    if num < 10:
        return chinese_digits[num]
    elif num == 10:
        return '十'
    elif num < 20:
        return '十' + chinese_digits[num % 10]
    elif num < 100:
        tens = num // 10
        ones = num % 10
        if ones == 0:
            return chinese_digits[tens] + '十'
        else:
            return chinese_digits[tens] + '十' + chinese_digits[ones]
    else:
        return str(num)


class HeaderNumbering:
    """标题编号管理器
    
    支持多级标题编号：
    - 一级标题：一、二、三、...（中文数字）
    - 二级标题：1、2、3、...（阿拉伯数字）
    - 三级标题：1）、2）、3）、...（阿拉伯数字+右括号）
    
    编号逻辑：同级标题需要连续编号，子标题在每个父级标题下重新开始编号
    """
    
    def __init__(self):
        self.level_counters = {}  # 存储每个级别的计数器
        self.last_level = 0  # 记录上一个标题的级别
        self.parent_levels = {}  # 记录每个级别的父级标题编号，用于判断是否遇到新的父级标题
    
    def enter_level(self, level: int) -> None:
        """进入指定级别的标题
        
        Args:
            level: 标题级别 (1-6)
        """
        # 一级标题需要连续编号，永远不重置
        # 二级标题在每个一级标题下重新开始编号，但同级需要连续编号
        # 三级及以下标题在每个父级标题下重新开始编号，但同级需要连续编号
        
        if level < self.last_level:
            # 如果当前级别小于上一个级别（从子标题跳回父级标题）
            # 需要判断是否遇到了新的更高层级标题
            # 如果父级标题编号没有变化，说明还在同一个父级标题下，应该继续累加
            # 如果父级标题编号变化了，说明遇到了新的更高层级标题，需要重置
            
            # 检查是否遇到了新的更高层级标题
            # 通过检查当前级别的父级标题编号是否变化来判断
            parent_level = level - 1
            current_parent_number = self.level_counters.get(parent_level, 0) if parent_level > 0 else 0
            last_parent_number = self.parent_levels.get(level, None)  # 使用None表示未记录过
            
            # 只有当之前记录过父级编号且父级编号发生变化时，才认为是遇到了新的更高层级标题
            if parent_level > 0 and last_parent_number is not None and current_parent_number != last_parent_number:
                # 遇到了新的更高层级标题，需要重置当前级别及其子级别的计数器
                for i in range(level, 7):
                    # 一级标题永远不重置（跳过level=1）
                    if i > 1:
                        self.level_counters[i] = 0
                # 重置后，当前级别应该从1开始编号
                self.level_counters[level] = 1
            else:
                # 还在同一个父级标题下，同级标题应该继续累加
                # 如果之前没有记录过，说明是第一次遇到这个级别，也应该累加
                self.level_counters[level] = self.level_counters.get(level, 0) + 1
            
            # 重置更深级别的计数器
            for i in range(level + 1, 7):
                self.level_counters[i] = 0
            
            # 更新父级标题编号记录
            if parent_level > 0:
                self.parent_levels[level] = current_parent_number
        elif level > self.last_level:
            # 如果当前级别更深，只重置更深级别的计数器
            for i in range(level + 1, 7):
                self.level_counters[i] = 0
            # 当前级别应该从1开始编号（在新的父级标题下）
            self.level_counters[level] = 1
            
            # 记录父级标题编号
            parent_level = level - 1
            if parent_level > 0:
                self.parent_levels[level] = self.level_counters.get(parent_level, 0)
        else:
            # 如果level == lastLevel，说明是同级标题，继续累加
            self.level_counters[level] = self.level_counters.get(level, 0) + 1
        
        # 更新上一个级别
        self.last_level = level
    
    def get_number(self, level: int) -> str:
        """获取当前编号，根据级别返回不同格式
        
        编号格式：
        - 一级标题：一、二、三、...
        - 二级标题：1、2、3、...（在当前一级标题下独立编号）
        - 三级标题：1）、2）、3）、...（在当前二级标题下独立编号）
        - 四级及以下：使用点号分隔的层级编号
        
        Args:
            level: 标题级别 (1-6)
        
        Returns:
            编号字符串
        """
        if level not in self.level_counters or self.level_counters[level] == 0:
            return ""
        
        # 根据标题级别决定格式
        if level == 1:
            # 一级标题：使用中文数字，如 "一、"
            num = self.level_counters.get(1, 0)
            return number_to_chinese(num) + "、" if num > 0 else ""
        elif level == 2:
            # 二级标题：使用阿拉伯数字，如 "1、"
            num = self.level_counters.get(2, 0)
            return str(num) + "、" if num > 0 else ""
        elif level == 3:
            # 三级标题：使用阿拉伯数字+右括号，如 "1）"
            num = self.level_counters.get(3, 0)
            return str(num) + "）" if num > 0 else ""
        elif level == 5:
            # 五级标题：使用阿拉伯数字+右括号，如 "1）"，在每个父级标题下重新开始编号
            num = self.level_counters.get(5, 0)
            return str(num) + "）" if num > 0 else ""
        elif level == 6:
            # 六级标题：使用阿拉伯数字+右括号，如 "1）"，在每个父级标题下重新开始编号
            num = self.level_counters.get(6, 0)
            return str(num) + "）" if num > 0 else ""
        else:
            # 四级：使用点号分隔的层级编号，如 "1.1.1"
            parts = []
            for i in range(1, level + 1):
                if i in self.level_counters and self.level_counters[i] > 0:
                    parts.append(str(self.level_counters[i]))
            return '.'.join(parts) + "、"


class DynamicWordDocumentCreator:
    """动态生成Word文档模板
    
    解析 Markdown 结构，创建 Word 文档框架
    插入占位符（${title}、${chart1}、${table1}）
    设置标题样式（自定义 Heading1-6 样式）
    实现标题自动编号
    设置段落格式（四号仿宋、1.5倍行距、首行缩进2字符）
    """
    
    @staticmethod
    def create_complete_template_from_markdown(file_path: str, markdown_content: str) -> None:
        """根据Markdown内容创建更完整的模板
        
        Args:
            file_path: 输出文件路径
            markdown_content: Markdown内容
        """
        document = Document()
        
        # 创建标题样式
        DynamicWordDocumentCreator._create_header_styles(document)
        
        # 创建标题段落
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        DynamicWordDocumentCreator._set_title_paragraph_style(title_paragraph)
        
        title_run = title_paragraph.add_run("${title}")
        title_run.bold = True
        title_run.font.size = Pt(16)  # 三号字体
        title_run.font.name = '仿宋'
        
        # 添加一个空行
        empty_paragraph = document.add_paragraph()
        DynamicWordDocumentCreator._set_default_paragraph_style(empty_paragraph)
        
        # 解析Markdown内容并创建相应的Word结构
        DynamicWordDocumentCreator._parse_and_create_document_structure(document, markdown_content)
        
        # 保存文档
        document.save(file_path)
    
    @staticmethod
    def _create_header_styles(document: Document) -> None:
        """创建自定义标题样式
        
        Args:
            document: Word文档对象
        """
        styles = document.styles
        
        # 创建标题1-6样式
        heading_configs = [
            ("Heading1", 1, 16),  # 一级标题：三号字体
            ("Heading2", 2, 16),  # 二级标题：三号字体
            ("Heading3", 3, 16),  # 三级标题：三号字体
            ("Heading4", 4, 16),
            ("Heading5", 5, 14),
            ("Heading6", 6, 12),
        ]
        
        for style_name, level, font_size in heading_configs:
            DynamicWordDocumentCreator._create_heading_style(styles, style_name, level, font_size)
    
    @staticmethod
    def _create_heading_style(styles, style_name: str, heading_level: int, font_size: int) -> None:
        """创建标题样式
        
        Args:
            styles: 样式集合
            style_name: 样式名称
            heading_level: 标题级别
            font_size: 字体大小
        """
        try:
            # 尝试获取现有样式 - 遍历查找而不是使用已弃用的ID查找
            existing_style = None
            for style in styles:
                if style.name == style_name:
                    existing_style = style
                    break

            if existing_style is None:
                # 创建新样式
                style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            else:
                style = existing_style
        except (KeyError, ValueError):
            # 如果样式不存在,创建新样式
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        
        # 设置字体
        font = style.font
        font.name = '仿宋'
        font.size = Pt(font_size)
        font.bold = True
        
        # 设置段落格式
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_before = Pt(12)  # 段前间距
        paragraph_format.space_after = Pt(6)  # 段后间距

        # 设置大纲级别（用于目录生成）
        paragraph_format.outline_level = heading_level - 1
    
    @staticmethod
    def _set_default_paragraph_style(paragraph) -> None:
        """设置默认段落样式 - 小四号仿宋，1.5倍行距，首行缩进2字符
        
        Args:
            paragraph: 段落对象
        """
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进2字符
        paragraph_format.space_before = Pt(0)  # 段前0磅
        paragraph_format.space_after = Pt(0)  # 段后0磅
    
    @staticmethod
    def _set_title_paragraph_style(paragraph) -> None:
        """设置标题段落样式
        
        Args:
            paragraph: 标题段落
        """
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_before = Pt(12)  # 标题前留一些空间
        paragraph_format.space_after = Pt(6)  # 标题后留一些空间
    
    @staticmethod
    def _parse_and_create_document_structure(document: Document, markdown_content: str) -> None:
        """解析Markdown内容并创建Word文档结构
        
        Args:
            document: Word文档对象
            markdown_content: Markdown内容
        """
        lines = markdown_content.split('\n')
        chart_index = 1
        table_index = 1
        image_index = 1

        # 初始化标题编号器
        header_numbering = HeaderNumbering()
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # 检查是否为标题
            header_info = MarkdownParser.is_header_line(line)
            if header_info:
                level, title = header_info
                
                # 更新标题编号
                header_numbering.enter_level(level)
                header_number = header_numbering.get_number(level)
                
                header_paragraph = document.add_paragraph()
                DynamicWordDocumentCreator._set_header_style(header_paragraph, level)
                
                header_run = header_paragraph.add_run(f"{header_number}{title}")
                header_run.bold = True
                header_run.font.name = '仿宋'

                # 根据标题级别设置字体大小
                # H1-H3: 三号(16pt), H4: 小四(12pt), H5-H6: 五号(10.5pt)
                font_sizes = {1: 16, 2: 16, 3: 16, 4: 12, 5: 10.5, 6: 10.5}
                header_run.font.size = Pt(font_sizes.get(level, 12))
                
                i += 1
                continue
            
            # 检查是否为ECharts图表
            if MarkdownParser.is_echarts_start_line(line):
                # 查找图表代码块的结束位置
                chart_code_lines = []
                i += 1  # 移动到下一行
                while i < len(lines) and not MarkdownParser.is_echarts_end_line(lines[i]):
                    chart_code_lines.append(lines[i])
                    i += 1
                
                chart_code = '\n'.join(chart_code_lines)
                
                # 创建图表占位符
                chart_title_paragraph = document.add_paragraph()
                chart_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                DynamicWordDocumentCreator._set_default_paragraph_style(chart_title_paragraph)
                
                chart_title_run = chart_title_paragraph.add_run(f"图表 {chart_index}：")
                chart_title_run.bold = True
                chart_title_run.font.name = '仿宋'
                chart_title_run.font.size = Pt(14)  # 四号字体
                
                # 创建图表占位符段落
                chart_paragraph = document.add_paragraph()
                chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                DynamicWordDocumentCreator._set_default_paragraph_style(chart_paragraph)
                
                chart_run = chart_paragraph.add_run(f"${{chart{chart_index}}}")
                
                chart_index += 1
                i += 1
                continue

            # 检查是否为图片
            image_info = MarkdownParser.is_image_line(line)
            if image_info:
                alt_text, image_url, title = image_info

                # 创建图片标题段落
                image_title_paragraph = document.add_paragraph()
                image_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                DynamicWordDocumentCreator._set_default_paragraph_style(image_title_paragraph)

                image_title_run = image_title_paragraph.add_run(f"图片 {image_index}：{title}")
                image_title_run.bold = True
                image_title_run.font.name = '仿宋'
                image_title_run.font.size = Pt(14)  # 四号字体

                # 创建图片占位符段落
                image_paragraph = document.add_paragraph()
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                DynamicWordDocumentCreator._set_default_paragraph_style(image_paragraph)

                image_run = image_paragraph.add_run(f"${{image{image_index}}}")

                image_index += 1
                i += 1
                continue

            # 检查是否为表格开始
            if MarkdownParser.is_table_start_line(line):
                # 收集表格的所有行
                table_lines = [line]
                i += 1  # 移动到下一行
                while i < len(lines) and (lines[i].startswith('|') or 
                                        re.match(r'^\|?\s*[-|:\s]+\|?\s*$', lines[i])):
                    table_lines.append(lines[i])
                    i += 1
                i -= 1  # 回退一行，因为循环会自动增加i
                
                # 创建表格占位符（不添加表格标题）
                table_paragraph = document.add_paragraph()
                table_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                DynamicWordDocumentCreator._set_default_paragraph_style(table_paragraph)
                
                table_run = table_paragraph.add_run(f"${{table{table_index}}}")
                
                table_index += 1
                i += 1
                continue
            
            # 普通段落
            if line.strip():
                paragraph = document.add_paragraph()
                DynamicWordDocumentCreator._set_default_paragraph_style(paragraph)
                
                # 格式化数字，添加千分符号
                formatted_line = format_number_with_thousands_separator(line)
                run = paragraph.add_run(formatted_line)
                run.font.name = '仿宋'
                run.font.size = Pt(12)  # 小四号字体
            
            i += 1
    
    @staticmethod
    def _set_header_style(paragraph, level: int) -> None:
        """设置标题样式
        
        Args:
            paragraph: 标题段落
            level: 标题级别
        """
        style_names = {
            1: "Heading1",
            2: "Heading2", 
            3: "Heading3",
            4: "Heading4",
            5: "Heading5",
            6: "Heading6"
        }
        
        style_name = style_names.get(level, "Heading1")
        try:
            paragraph.style = style_name
        except KeyError:
            # 如果样式不存在，使用默认样式
            pass
