"""
Word 文档生成器
对应 Java 版本的 PoiWordGenerator.java
"""

import re
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from ..models.word_params import WordParams, TextParam, TableParam, ImageParam
from ..models.chart_table import ChartTable


def format_number_with_thousands_separator(text: str) -> str:
    """格式化文本中的数字，添加千分符号
    
    例如：1000 -> 1,000，1234567 -> 1,234,567
    
    Args:
        text: 输入文本
        
    Returns:
        格式化后的文本
    """
    # 匹配整数和小数的正则表达式
    # 匹配整数：\d+（一个或多个数字）
    # 匹配小数：\d+\.\d+（数字.数字）
    def replace_number(match):
        number_str = match.group(0)
        # 如果是小数，分别处理整数部分和小数部分
        if '.' in number_str:
            integer_part, decimal_part = number_str.split('.')
            # 格式化整数部分
            formatted_integer = format(int(integer_part), ',')
            return f"{formatted_integer}.{decimal_part}"
        else:
            # 格式化整数
            return format(int(number_str), ',')
    
    # 匹配数字（包括整数和小数）
    # \d+ 匹配一个或多个数字
    # (\.\d+)? 可选的小数部分
    pattern = r'\d+(\.\d+)?'
    return re.sub(pattern, replace_number, text)


class PoiWordGenerator:
    """Word 文档生成器
    
    读取模板文档，替换文本占位符，插入表格（带表头背景色），插入图表
    保存最终文档
    """
    
    @staticmethod
    def build_doc(params: WordParams, template_file: str, output_file: str) -> bool:
        """构建 Word 文档
        
        Args:
            params: Word 参数对象
            template_file: 模板文件路径
            output_file: 输出文件路径
            
        Returns:
            是否成功构建文档
        """
        try:
            # 打开模板文档
            document = Document(template_file)
            
            # 替换段落中的占位符
            PoiWordGenerator._replace_paragraphs(document, params)
            
            # 替换图表占位符
            PoiWordGenerator._replace_charts(document, params)
            
            # 保存文档
            document.save(output_file)
            return True
            
        except Exception as e:
            print(f"构建文档时出错: {e}")
            return False
    
    @staticmethod
    def _replace_paragraphs(document: Document, params: WordParams) -> None:
        """替换段落中的占位符
        
        Args:
            document: Word 文档对象
            params: Word 参数对象
        """
        for paragraph in document.paragraphs:
            PoiWordGenerator._replace_paragraph_content(paragraph, params)
    
    @staticmethod
    def _replace_paragraph_content(paragraph, params: WordParams) -> None:
        """替换段落内容中的占位符
        
        Args:
            paragraph: 段落对象
            params: Word 参数对象
        """
        # 获取段落的所有文本
        full_text = paragraph.text
        
        # 查找所有占位符
        placeholder_pattern = r'\$\{([^}]+)\}'
        matches = list(re.finditer(placeholder_pattern, full_text))
        
        if not matches:
            return
        
        # 从后往前替换，避免位置偏移
        for match in reversed(matches):
            placeholder = match.group(0)  # ${key}
            key = match.group(1)  # key
            
            # 获取参数值
            param = params.get_param(key)
            
            if param is None:
                continue
            
            if isinstance(param, TextParam):
                # 替换文本
                PoiWordGenerator._replace_text_in_paragraph(paragraph, placeholder, param.msg)
            elif isinstance(param, TableParam):
                # 替换为表格
                PoiWordGenerator._replace_with_table(paragraph, placeholder, param.data)
            elif isinstance(param, ImageParam):
                # 替换为图像
                PoiWordGenerator._replace_with_image(paragraph, placeholder, param)
    
    @staticmethod
    def _replace_text_in_paragraph(paragraph, placeholder: str, replacement: str) -> None:
        """在段落中替换文本
        
        Args:
            paragraph: 段落对象
            placeholder: 占位符文本
            replacement: 替换文本
        """
        # 清空段落内容
        paragraph.clear()
        
        # 格式化数字，添加千分符号
        formatted_text = format_number_with_thousands_separator(replacement)
        
        # 添加新文本
        run = paragraph.add_run(formatted_text)
        run.font.name = '仿宋'
        run.font.size = Pt(14)  # 四号字体
    
    @staticmethod
    def _replace_with_table(paragraph, placeholder: str, table_data: List[List[str]]) -> None:
        """用表格替换占位符
        
        Args:
            paragraph: 段落对象
            placeholder: 占位符文本
            table_data: 表格数据
        """
        # 清空段落内容
        paragraph.clear()
        
        # 获取文档对象
        document = paragraph._element.getparent().getparent()
        
        # 找到段落在文档中的位置
        paragraph_element = paragraph._element
        parent_element = paragraph_element.getparent()
        
        # 创建表格
        table_element = OxmlElement('w:tbl')
        
        # 创建表格属性
        tblPr = OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        
        # 设置边框
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        table_element.append(tblPr)
        
        # 创建表格网格
        tblGrid = OxmlElement('w:tblGrid')
        for i in range(len(table_data[0]) if table_data else 0):
            gridCol = OxmlElement('w:gridCol')
            tblGrid.append(gridCol)
        table_element.append(tblGrid)
        
        # 添加表格行
        for i, row_data in enumerate(table_data):
            tr = OxmlElement('w:tr')
            
            for j, cell_data in enumerate(row_data):
                tc = OxmlElement('w:tc')
                
                # 设置单元格属性
                tcPr = OxmlElement('w:tcPr')
                tcBorders = OxmlElement('w:tcBorders')
                
                # 设置单元格边框
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                
                tcPr.append(tcBorders)
                
                # 设置表头背景色
                if i == 0:  # 第一行
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:val'), 'clear')
                    shading.set(qn('w:color'), 'auto')
                    shading.set(qn('w:fill'), 'B4C6E7')
                    tcPr.append(shading)
                
                tc.append(tcPr)
                
                # 添加段落和文本
                p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                pPr.append(jc)
                p.append(pPr)
                
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                font = OxmlElement('w:rFonts')
                font.set(qn('w:ascii'), '仿宋')
                font.set(qn('w:hAnsi'), '仿宋')
                font.set(qn('w:eastAsia'), '仿宋')
                rPr.append(font)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '22')  # 11pt (11号字体)
                rPr.append(sz)
                r.append(rPr)
                
                t = OxmlElement('w:t')
                # 格式化数字，添加千分符号
                t.text = format_number_with_thousands_separator(cell_data)
                r.append(t)
                p.append(r)
                
                tc.append(p)
                tr.append(tc)
            
            table_element.append(tr)
        
        # 在段落后插入表格
        parent_element.insert(parent_element.index(paragraph_element) + 1, table_element)
    
    @staticmethod
    def _replace_with_image(paragraph, placeholder: str, image_param: ImageParam) -> None:
        """用图像替换占位符

        Args:
            paragraph: 段落对象
            placeholder: 占位符文本
            image_param: 图像参数
        """
        # 清空段落内容
        paragraph.clear()

        # 设置段落居中对齐
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加图像
        try:
            from io import BytesIO
            
            # 验证图片数据
            if not image_param.image_data or len(image_param.image_data) == 0:
                paragraph.add_run(f"[图像数据为空]")
                return
            
            # 创建图片流并确保位置在开头
            image_stream = BytesIO(image_param.image_data)
            image_stream.seek(0)  # 重置流位置到开头，确保 add_picture 能正确读取

            # 计算图片尺寸（使用 Inches，Word 的默认单位）
            # 使用更合理的 DPI 转换（Word 默认 96 DPI，但实际显示可能不同）
            # 为了更好的显示效果，使用 72 DPI 进行计算（标准屏幕 DPI）
            dpi = 72  # 使用标准 DPI
            width_inches = Inches(image_param.width / dpi)
            height_inches = Inches(image_param.height / dpi)

            # 限制最大尺寸（适应 Word A4 页面，留出边距）
            # A4 页面宽度约 8.27 英寸，高度约 11.69 英寸
            # 考虑页边距，实际可用宽度约 6.5 英寸，高度约 9.5 英寸
            max_width = Inches(6.5)  # 最大 6.5 英寸
            max_height = Inches(9.5)  # 最大 9.5 英寸

            # 等比例缩放，确保图片完全显示
            if width_inches > max_width or height_inches > max_height:
                width_ratio = max_width / width_inches if width_inches > max_width else 1
                height_ratio = max_height / height_inches if height_inches > max_height else 1
                ratio = min(width_ratio, height_ratio)
                width_inches = width_inches * ratio
                height_inches = height_inches * ratio

            # 确保最小尺寸（避免图片过小）
            min_width = Inches(1)  # 最小 1 英寸
            min_height = Inches(0.75)  # 最小 0.75 英寸
            
            if width_inches < min_width:
                # 等比例调整
                height_inches = height_inches * (min_width / width_inches)
                width_inches = min_width
            
            if height_inches < min_height:
                # 等比例调整
                width_inches = width_inches * (min_height / height_inches)
                height_inches = min_height

            # 添加图片到段落
            paragraph.add_run().add_picture(image_stream, width=width_inches, height=height_inches)

        except Exception as e:
            print(f"添加图像时出错: {e}")
            print(f"图片尺寸: {image_param.width}x{image_param.height}, 数据大小: {len(image_param.image_data) if image_param.image_data else 0} 字节")
            paragraph.add_run(f"[图像加载失败: {str(e)}]")
            import traceback
            traceback.print_exc()
    
    @staticmethod
    def _set_cell_background(cell, color: str) -> None:
        """设置单元格背景色
        
        Args:
            cell: 单元格对象
            color: 颜色值（如 "B4C6E7"）
        """
        try:
            # 获取单元格的 XML 元素
            cell_xml = cell._tc
            
            # 创建背景色元素
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), color)
            
            # 获取或创建段落属性
            p_pr = cell.paragraphs[0]._element.get_or_add_pPr()
            p_pr.append(shading)
            
        except Exception as e:
            print(f"设置单元格背景色时出错: {e}")
    
    @staticmethod
    def _replace_charts(document: Document, params: WordParams) -> None:
        """替换图表占位符
        
        Args:
            document: Word 文档对象
            params: Word 参数对象
        """
        # 查找所有图表占位符
        for paragraph in document.paragraphs:
            text = paragraph.text
            chart_pattern = r'\$\{chart(\d+)\}'
            matches = re.finditer(chart_pattern, text)
            
            for match in matches:
                chart_key = f"chart{match.group(1)}"
                chart_table = params.get_chart(chart_key)
                
                if chart_table:
                    # 清空段落内容
                    paragraph.clear()
                    
                    # 创建图表
                    try:
                        PoiWordGenerator._create_chart_in_paragraph(paragraph, chart_table)
                    except Exception as e:
                        print(f"创建图表时出错: {e}")
                        paragraph.add_run(f"图表 {chart_key} 创建失败")
    
    @staticmethod
    def _create_chart_in_paragraph(paragraph, chart_table: ChartTable) -> None:
        """在段落中创建图表

        Args:
            paragraph: 段落对象
            chart_table: 图表数据表
        """
        # 准备表格数据
        x_data = list(chart_table.get_x_axis())
        y_axes = chart_table.get_y_axis_dict()

        if not x_data or not y_axes:
            paragraph.add_run("图表数据为空")
            return

        # 在段落中添加图表标题
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = paragraph.add_run(f"【图表】{chart_table.title}")
        title_run.bold = True
        title_run.font.name = '仿宋'
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = None  # 使用默认颜色

        # 创建表格数据
        table_data = []

        # 添加表头
        header_row = ["类别"]
        for series_name in y_axes.keys():
            header_row.append(series_name)
        table_data.append(header_row)

        # 添加数据行
        for i, x_value in enumerate(x_data):
            row = [str(x_value)]
            for series_name, y_axis in y_axes.items():
                if i < len(y_axis.data_list):
                    row.append(str(y_axis.data_list[i]))
                else:
                    row.append("")
            table_data.append(row)

        # 使用XML方式创建表格
        paragraph_element = paragraph._element
        parent_element = paragraph_element.getparent()

        # 创建表格元素
        table_element = OxmlElement('w:tbl')

        # 创建表格属性
        tblPr = OxmlElement('w:tblPr')

        # 设置表格居中对齐
        tblJc = OxmlElement('w:jc')
        tblJc.set(qn('w:val'), 'center')
        tblPr.append(tblJc)

        # 设置表格边框
        tblBorders = OxmlElement('w:tblBorders')

        # 设置边框
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')  # 加粗边框
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '4472C4')  # 蓝色边框
            tblBorders.append(border)

        tblPr.append(tblBorders)
        table_element.append(tblPr)

        # 创建表格网格
        tblGrid = OxmlElement('w:tblGrid')
        for i in range(len(table_data[0]) if table_data else 0):
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), '2000')  # 设置列宽
            tblGrid.append(gridCol)
        table_element.append(tblGrid)

        # 添加表格行
        for i, row_data in enumerate(table_data):
            tr = OxmlElement('w:tr')

            # 设置行高
            trPr = OxmlElement('w:trPr')
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '400')
            trPr.append(trHeight)
            tr.append(trPr)

            for cell_data in row_data:
                tc = OxmlElement('w:tc')

                # 设置单元格属性
                tcPr = OxmlElement('w:tcPr')

                # 设置单元格垂直对齐
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

                # 设置表头背景色 - 使用更鲜明的颜色
                if i == 0:  # 第一行（表头）
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:val'), 'clear')
                    shading.set(qn('w:color'), 'auto')
                    shading.set(qn('w:fill'), '4472C4')  # 深蓝色
                    tcPr.append(shading)
                elif i % 2 == 1:  # 奇数行使用浅色背景
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:val'), 'clear')
                    shading.set(qn('w:color'), 'auto')
                    shading.set(qn('w:fill'), 'E7E6E6')  # 浅灰色
                    tcPr.append(shading)

                tc.append(tcPr)

                # 添加段落和文本
                p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                pPr.append(jc)
                p.append(pPr)

                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')

                # 设置字体
                font = OxmlElement('w:rFonts')
                font.set(qn('w:ascii'), '仿宋')
                font.set(qn('w:hAnsi'), '仿宋')
                font.set(qn('w:eastAsia'), '仿宋')
                rPr.append(font)

                # 设置字号 - 统一使用11号字体
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '22')  # 11pt (11号字体)
                rPr.append(sz)

                # 表头使用白色加粗字体
                if i == 0:
                    bold = OxmlElement('w:b')
                    rPr.append(bold)
                    color = OxmlElement('w:color')
                    color.set(qn('w:val'), 'FFFFFF')  # 白色
                    rPr.append(color)

                r.append(rPr)

                t = OxmlElement('w:t')
                # 格式化数字，添加千分符号
                t.text = format_number_with_thousands_separator(cell_data)
                r.append(t)
                p.append(r)

                tc.append(p)
                tr.append(tc)

            table_element.append(tr)

        # 在段落后插入表格
        parent_element.insert(parent_element.index(paragraph_element) + 1, table_element)
